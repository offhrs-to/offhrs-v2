"""
Generates docs/OFFHRS_TEST_PLAN.docx — a printable / shareable test plan
with Word-compatible interactive checkbox content controls.

Usage:
    python scripts/generate_test_plan.py
"""
from __future__ import annotations

import os
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

OUTPUT_PATH = os.path.join("docs", "OFFHRS_TEST_PLAN.docx")

# Sections: title, optional intro paragraph(s), and a list of (subhead or None, [checks]).
SECTIONS: list[dict] = [
    {
        "number": "0",
        "title": "Pre-Test Setup",
        "intro": [
            "Use separate test identities so account deletion does not confuse results.",
        ],
        "groups": [
            (
                "Test identities",
                [
                    "Create a consumer-only test account (e.g. consumer-test@…)",
                    "Create a vendor-only test account (e.g. vendor-test@…)",
                    "Create a dual-role account used for both mobile consumer and partner vendor",
                    "Choose the Google OAuth account to verify account-picker behaviour",
                ],
            ),
            (
                "Stripe sandbox readiness",
                [
                    "Lite price is $29 CAD/month, recurring monthly",
                    "Pro price is $49 CAD/month, recurring monthly",
                    "Both prices set to “exclusive of tax”",
                    "Lite and Pro products have a SaaS tax code (e.g. txcd_10103000)",
                    "Stripe Tax registration for Canada / Ontario exists in the sandbox if HST is expected",
                ],
            ),
        ],
    },
    {
        "number": "1",
        "title": "Terms & Policies",
        "intro": [
            "Verify the new Fresha-style Terms overview and the 5 detail pages.",
        ],
        "groups": [
            (
                "Terms overview page (/terms)",
                [
                    "Header shows offhrs logo, “For Partners”, and “Contact us”",
                    "Header does NOT include a “Workshops” link",
                    "Tab nav row shows Overview, Terms of Use, Privacy Policy, Service Terms, Data Protection Addendum, Cookie Policy",
                    "Five cards render with correct titles and summaries",
                    "“Last updated” date is displayed",
                    "Footer shows For Business and Legal columns",
                ],
            ),
            (
                "Detail page navigation",
                [
                    "/terms/terms-of-use loads with title “Terms of Use”",
                    "/terms/privacy-policy loads with title “Privacy Policy”",
                    "/terms/service-terms loads with title “Service Terms”",
                    "/terms/data-protection loads with title “Data Protection Addendum”",
                    "/terms/cookies loads with title “Cookie Policy”",
                    "Each detail page has a back link to /terms",
                    "Each detail page lists related policy links at the bottom",
                ],
            ),
            (
                "Backwards compatibility",
                [
                    "/privacy redirects to /terms/privacy-policy",
                    "Web footer Terms link still works",
                    "Partner dashboard sidebar shows “Terms & policies” above Sign out",
                    "Clicking “Terms & policies” in partner sidebar opens /terms in a new tab",
                ],
            ),
        ],
    },
    {
        "number": "2",
        "title": "Mobile Policy Links",
        "intro": [
            "Install the latest preview build (or wait for OTA) before running this section.",
        ],
        "groups": [
            (
                "Signed-out Profile / Sign-In screen",
                [
                    "Sign-in area says “By continuing you agree to our Terms & policies.”",
                    "Tapping the link opens the web Terms Overview page",
                ],
            ),
            (
                "Signed-in Profile screen",
                [
                    "Bottom of Profile shows a single “Terms & policies” link (no Privacy / Terms / Listing disclaimer trio)",
                    "Tapping the link opens /terms",
                ],
            ),
            (
                "Workshops disclaimer footer",
                [
                    "Disclaimer footer shows a single “Terms & policies” link",
                    "Tapping the link opens /terms",
                ],
            ),
        ],
    },
    {
        "number": "3",
        "title": "Google OAuth Account Picker",
        "intro": [
            "Validates the fix where Google previously auto-logged into a remembered account.",
        ],
        "groups": [
            (
                "Account selection",
                [
                    "After tapping Continue with Google, the Google account chooser is shown",
                    "Google does NOT silently re-use the previously signed-in account",
                    "Signing out and tapping Continue with Google again still shows the chooser",
                    "Selecting a different Google account logs into that account in the app",
                ],
            ),
        ],
    },
    {
        "number": "4",
        "title": "Consumer Account Deletion",
        "intro": [
            "Mobile “Delete my account” deletes only the consumer role and preserves any vendor login on the same email.",
        ],
        "groups": [
            (
                "Consumer-only account",
                [
                    "Mobile signs out after delete",
                    "The deleted account cannot log back in",
                    "profiles row is gone in Supabase",
                    "bookings, user_event_saves, user_vendor_saves, vendor_reviews, profile_category_experience rows are gone",
                    "Affected event slot counts are reconciled",
                ],
            ),
            (
                "Dual-role account (mobile delete)",
                [
                    "Consumer profile and consumer-owned data are deleted",
                    "vendor_profiles row remains intact",
                    "auth.users row remains intact",
                    "Same email can still sign into the partner dashboard",
                    "Partner dashboard loads correctly with vendor data intact",
                ],
            ),
        ],
    },
    {
        "number": "5",
        "title": "Vendor Account Deletion",
        "intro": [
            "Partner Settings → Danger zone → Delete vendor account hits /api/partners/account/delete.",
        ],
        "groups": [
            (
                "Vendor-only account",
                [
                    "Confirm dialog spells out subscription cancel + permanent data loss",
                    "Stripe subscription is canceled (status = canceled)",
                    "vendor_profiles row is deleted",
                    "events, bookings (host), vendor_subscriptions, vendor_payouts, vendor_calendar_connections, vendor_reviews are deleted (cascade)",
                    "Vendor-only login: auth.users row is also deleted",
                    "Vendor-only login: Email can be re-used for a fresh signup",
                ],
            ),
            (
                "Dual-role account (partner delete)",
                [
                    "vendor_profiles row and vendor data are deleted",
                    "profiles row remains intact",
                    "auth.users row remains intact",
                    "Mobile consumer login still works for the same email",
                    "Browser session is NOT signed out (consumer profile still present)",
                ],
            ),
        ],
    },
    {
        "number": "6",
        "title": "Partner Signup Flow",
        "intro": [
            "Test both Lite and Pro, plus the existing-customer retry case that broke earlier.",
        ],
        "groups": [
            (
                "Lite plan signup",
                [
                    "/partners/signup loads",
                    "Account setup completes",
                    "Selecting Lite shows $29 CAD/month",
                    "Add payment & start trial redirects to Stripe Checkout without errors",
                    "Stripe Checkout completes with a test card",
                    "User lands in the partner dashboard",
                    "vendor_profiles.stripe_customer_id is set",
                    "vendor_subscriptions row exists with subscription_tier = lite",
                    "Stripe shows a 30-day trial",
                ],
            ),
            (
                "Pro plan signup",
                [
                    "Selecting Pro shows $49 CAD/month",
                    "Stripe Checkout completes with Pro price",
                    "vendor_subscriptions.subscription_tier = pro",
                ],
            ),
            (
                "Existing customer retry",
                [
                    "Cancel mid-checkout, then click Add payment & start trial again",
                    "No “Tax ID collection requires updating business name on the customer” error appears",
                    "Stripe Checkout opens successfully on retry",
                ],
            ),
        ],
    },
    {
        "number": "7",
        "title": "Stripe Tax for Vendor Subscription",
        "intro": [
            "Account-level Stripe Tax must be configured for HST to actually be charged.",
        ],
        "groups": [
            (
                "Tax wiring",
                [
                    "Stripe sandbox has a Canada / Ontario tax registration",
                    "Lite + Pro prices use “Exclusive of tax”",
                    "Lite + Pro products have a SaaS tax code",
                    "Lite upcoming invoice shows tax line (≈ $3.77 HST on $29)",
                    "Pro upcoming invoice shows tax line (≈ $6.37 HST on $49)",
                    "If tax is $0, the Stripe tax-calculation-details message identifies the missing piece",
                ],
            ),
        ],
    },
    {
        "number": "8",
        "title": "Partner Dashboard Core Pages",
        "intro": [
            "Run as a vendor with an active or trial subscription and at least one workshop.",
        ],
        "groups": [
            (
                "Overview",
                [
                    "KPIs load without console errors",
                    "Refunded / cancelled bookings are NOT counted as active revenue",
                    "Recent bookings list shows confirmed and refunded states",
                    "Activity graph reflects refunds as churn",
                ],
            ),
            (
                "Workshops",
                [
                    "New workshop appears in Workshops page",
                    "Capacity displays correctly after booking",
                    "Capacity restores after refund",
                    "Capacity reconciles after consumer account deletion",
                ],
            ),
            (
                "Bookings",
                [
                    "Confirmed booking appears after consumer checkout",
                    "Refunded booking shows status Refunded",
                    "Refunded price is visually de-emphasized (strikethrough / muted)",
                ],
            ),
            (
                "Clients",
                [
                    "Consumer with name / email appears in Clients",
                    "Consumer phone (if provided) appears in Clients",
                ],
            ),
            (
                "Payouts",
                [
                    "Successful booking contributes to payouts / revenue numbers",
                    "Refunded bookings do not inflate totals",
                ],
            ),
        ],
    },
    {
        "number": "9",
        "title": "Workshop Creation Limits",
        "intro": [
            "Lite is capped to 4 new workshop sessions per Stripe billing period; Pro is unlimited.",
        ],
        "groups": [
            (
                "Lite cap",
                [
                    "Lite vendor can create 4 sessions in the current billing period",
                    "5th creation attempt is blocked with an upgrade or limit message",
                ],
            ),
            (
                "Pro unlimited",
                [
                    "Pro vendor can create more than 4 sessions without hitting the cap",
                ],
            ),
        ],
    },
    {
        "number": "10",
        "title": "Calendar Sync (Google + Outlook)",
        "intro": [
            "First-party OAuth integration (no Cal.com). Verifies the cleanup we just did.",
        ],
        "groups": [
            (
                "Google Calendar",
                [
                    "Connect Google Calendar from Partner Dashboard → Calendar",
                    "OAuth completes and returns with calendar_connected=google in the URL",
                    "Success banner reads “Connected Google Calendar…”",
                    "vendor_calendar_connections row exists with provider = google",
                    "Publishing a workshop creates the event on Google Calendar",
                    "events.google_calendar_event_id is populated",
                    "Editing the workshop updates the Google Calendar event",
                    "Unpublishing / deleting the workshop removes the Google Calendar event",
                ],
            ),
            (
                "Microsoft Outlook",
                [
                    "Connect Outlook returns with calendar_connected=microsoft",
                    "Success banner reads “Connected Outlook…”",
                    "events.microsoft_outlook_event_id is populated for published workshops",
                ],
            ),
            (
                "Disconnect",
                [
                    "Disconnecting removes the vendor_calendar_connections row",
                    "Events offhrs added are removed from the external calendar",
                ],
            ),
        ],
    },
    {
        "number": "11",
        "title": "Consumer Booking Flow",
        "intro": [
            "Mobile and web checkout. Validates tax columns, slot decrement, and confirmation email.",
        ],
        "groups": [
            (
                "Successful booking",
                [
                    "Payment succeeds with Stripe test card",
                    "bookings row is created with status confirmed",
                    "user_id is set when logged in",
                    "vendor_id and event_id are set",
                    "Tax columns are populated if tax applies",
                    "Confirmation email is received",
                    "Available slots decrease by quantity booked",
                    "Partner dashboard Bookings + Overview show the booking",
                ],
            ),
            (
                "Multiple tickets on same event",
                [
                    "Same user can book the same workshop a second time (no UNIQUE constraint error)",
                    "Capacity decrements correctly across multiple bookings",
                ],
            ),
            (
                "Guest / web booking",
                [
                    "Logged-out web booking succeeds when allowed",
                    "bookings row stores guest name + email",
                    "Partner dashboard still sees the guest booking",
                ],
            ),
        ],
    },
    {
        "number": "12",
        "title": "Refund & Cancellation Flow",
        "intro": [
            "Major area we fixed: status sync, slot restoration, dashboard reflection.",
        ],
        "groups": [
            (
                "Self-cancel outside refund window",
                [
                    "Stripe refund is created against the original PaymentIntent",
                    "Refund email is sent to the consumer",
                    "bookings.status becomes refunded and refunded_at is set",
                    "Partner dashboard Bookings shows Refunded",
                    "Activity chart counts the refund as churn",
                    "Workshop capacity restores",
                ],
            ),
            (
                "Inside refund window",
                [
                    "Refund is blocked or surfaces the vendor’s policy",
                    "Booking remains active",
                    "Slot counts do not change",
                ],
            ),
            (
                "Partner-issued refund",
                [
                    "Refund button on the partner dashboard issues a Stripe refund",
                    "Booking status becomes refunded",
                    "Capacity restores",
                    "Metrics update",
                ],
            ),
            (
                "Double refund protection",
                [
                    "Attempting a second refund on an already-refunded booking does not create a duplicate Stripe refund",
                    "App handles the “already refunded” case gracefully",
                ],
            ),
        ],
    },
    {
        "number": "13",
        "title": "Slot Reconciliation",
        "intro": [
            "Self-healing logic that runs on dashboard load and after deletions / refunds.",
        ],
        "groups": [
            (
                "After refund",
                [
                    "Workshop with 6 capacity → book 2 → refund 1 → dashboard shows 5/6 free",
                ],
            ),
            (
                "After consumer account deletion",
                [
                    "Deleting a consumer that had bookings restores their seats on the partner dashboard",
                ],
            ),
            (
                "On dashboard load",
                [
                    "Loading /partners/dashboard reconciles any drift in events.available_slots",
                    "Loading /partners/dashboard/sessions reconciles any drift",
                    "Loading /partners/dashboard/bookings reconciles any drift",
                ],
            ),
        ],
    },
    {
        "number": "14",
        "title": "Mobile OTA Update",
        "intro": [
            "Validates the cold-start update logic we added so preview builds pick up JS bundle changes.",
        ],
        "groups": [
            (
                "Update cycle",
                [
                    "Publish an Expo OTA update to the preview channel",
                    "Fully close the preview app",
                    "Reopen the preview app",
                    "App fetches and reloads automatically without a manual rebuild",
                ],
            ),
            (
                "Latest features visible after OTA",
                [
                    "Mobile Profile shows the single “Terms & policies” link",
                    "Google OAuth shows the account picker every time",
                    "Account deletion error messages include the failing stage from the backend",
                ],
            ),
        ],
    },
    {
        "number": "15",
        "title": "Auth / Role Separation Matrix",
        "intro": [
            "Confirms the three account states behave correctly end-to-end.",
        ],
        "groups": [
            (
                "Consumer-only login",
                [
                    "Mobile login works",
                    "Partner dashboard redirects to /partners/signup",
                    "Mobile delete removes auth.users",
                ],
            ),
            (
                "Vendor-only login",
                [
                    "Partner dashboard works",
                    "Partner delete removes auth.users when no consumer profile exists",
                ],
            ),
            (
                "Dual-role login",
                [
                    "Mobile app loads consumer experience",
                    "Partner dashboard loads vendor experience",
                    "Mobile delete preserves vendor and auth.users",
                    "Partner delete preserves consumer and auth.users",
                    "Each role’s data is cleanly separated after the other is deleted",
                ],
            ),
        ],
    },
    {
        "number": "16",
        "title": "Final Regression Smoke",
        "intro": [
            "Run after all critical paths above are green.",
        ],
        "groups": [
            (
                "Marketing + auth surfaces",
                [
                    "offhrs.app home loads",
                    "/terms loads",
                    "/privacy redirects",
                    "/partners loads",
                    "/partners/signup works",
                    "/partners/login works",
                ],
            ),
            (
                "Authenticated surfaces",
                [
                    "/partners/dashboard loads for active / trialing vendor",
                    "Mobile app launches and workshop list loads",
                    "Mobile booking checkout works end-to-end",
                ],
            ),
            (
                "Vercel function health",
                [
                    "No 500s on /api/book",
                    "No 500s on /api/book/confirm",
                    "No 500s on /api/account/delete",
                    "No 500s on /api/partners/account/delete",
                    "No 500s on /api/partners/checkout",
                    "No 500s on /api/partners/sessions",
                    "No 500s on /api/partners/bookings",
                    "Stripe webhooks (charge.refunded, customer.subscription.*) show 2xx",
                ],
            ),
        ],
    },
    {
        "number": "★",
        "title": "30-Minute Smoke Path",
        "intro": [
            "Use this when you only have time for the highest-value end-to-end loop.",
        ],
        "groups": [
            (
                "Critical loop",
                [
                    "Sign up as a fresh vendor on Lite → Stripe Checkout → dashboard loads",
                    "Create a workshop from the dashboard",
                    "Book that workshop from the mobile app",
                    "Confirm partner dashboard capacity and Bookings list update",
                    "Refund the booking from the mobile app",
                    "Confirm dashboard reflects Refunded status and capacity restores",
                    "Delete the consumer account from mobile",
                    "Confirm the same email’s vendor dashboard still exists (if dual-role)",
                    "Delete the vendor account from partner Settings",
                    "Confirm vendor data is gone in Supabase and Stripe subscription is canceled",
                    "Open /terms on web and tap “Terms & policies” in the mobile app",
                ],
            ),
        ],
    },
]


def add_checkbox(paragraph, label: str) -> None:
    """Add a Word interactive checkbox (sdt content control) + label text."""
    sdt = OxmlElement("w:sdt")

    sdt_pr = OxmlElement("w:sdtPr")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Segoe UI Symbol")
    rfonts.set(qn("w:hAnsi"), "Segoe UI Symbol")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rpr.append(sz)
    sdt_pr.append(rpr)

    # w14:checkbox container with checked/unchecked symbols.
    nsmap_w14 = "http://schemas.microsoft.com/office/word/2010/wordml"
    checkbox = OxmlElement("w14:checkbox")
    # Need to register the w14 namespace on the element manually.
    checkbox.set(qn("w14:checked") if False else "{%s}checked" % nsmap_w14, "0")
    checked = OxmlElement("w14:checkedState")
    checked.set("{%s}val" % nsmap_w14, "2612")
    checked.set("{%s}font" % nsmap_w14, "MS Gothic")
    unchecked = OxmlElement("w14:uncheckedState")
    unchecked.set("{%s}val" % nsmap_w14, "2610")
    unchecked.set("{%s}font" % nsmap_w14, "MS Gothic")
    checkbox.append(checked)
    checkbox.append(unchecked)
    sdt_pr.append(checkbox)
    sdt.append(sdt_pr)

    sdt_content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    rpr2 = deepcopy(rpr)
    run.append(rpr2)
    t = OxmlElement("w:t")
    t.text = "\u2610"  # ☐
    run.append(t)
    sdt_content.append(run)
    sdt.append(sdt_content)

    paragraph._p.append(sdt)

    # Label text after the checkbox.
    label_run = paragraph.add_run("  " + label)
    label_run.font.size = Pt(11)
    label_run.font.name = "Calibri"


def style_heading(p, text: str, level: int) -> None:
    p.style = p.part.document.styles[f"Heading {level}"]
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def build() -> None:
    doc = Document()

    # Tighten margins a bit.
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Title.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("offhrs — Production-Readiness Test Plan")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x5D, 0x75, 0x5D)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(
        "Verifies the Terms overhaul, Cal.com cleanup, vendor / consumer "
        "delete separation, partner signup fixes, and Stripe Tax wiring."
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        "Run order: Vercel preview + Supabase staging + Stripe sandbox first, "
        "then repeat the critical paths in production after deploy."
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    for section in SECTIONS:
        # Section heading.
        h = doc.add_paragraph()
        run = h.add_run(f"{section['number']}. {section['title']}")
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        # Coloured bar via paragraph border-bottom would require XML; skipping for portability.

        for line in section.get("intro", []):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            r.italic = True

        for group_title, checks in section["groups"]:
            if group_title:
                gh = doc.add_paragraph()
                gh_run = gh.add_run(group_title)
                gh_run.bold = True
                gh_run.font.size = Pt(11.5)
                gh_run.font.color.rgb = RGBColor(0x5D, 0x75, 0x5D)
                gh.paragraph_format.space_before = Pt(8)
                gh.paragraph_format.space_after = Pt(2)

            for check in checks:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_after = Pt(2)
                add_checkbox(p, check)

        doc.add_paragraph()  # spacer between sections

    # Footer line.
    foot = doc.add_paragraph()
    foot_run = foot.add_run(
        "Generated from scripts/generate_test_plan.py — regenerate after each "
        "release to refresh this checklist."
    )
    foot_run.italic = True
    foot_run.font.size = Pt(9)
    foot_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
